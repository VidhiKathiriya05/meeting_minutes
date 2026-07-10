from pathlib import Path
import json
from docx import Document
from docx.shared import Pt

OUTPUT_DIR = Path("output")
OUTPUT_DIR.mkdir(exist_ok=True)


def load_json(value):
    if not value:
        return []

    if isinstance(value, list):
        return value

    try:
        return json.loads(value)
    except:
        return []


def add_heading(doc, text):
    heading = doc.add_heading(text, level=2)
    heading.runs[0].font.size = Pt(14)


def add_list(doc, items):

    if not items:
        return

    for item in items:

        if isinstance(item, dict):

            task = item.get("task", "")
            owner = item.get("owner", "")
            deadline = item.get("deadline", "")

            p = doc.add_paragraph(style="List Bullet")
            p.add_run(task).bold = True

            if owner:
                doc.add_paragraph(f"Owner: {owner}")

            if deadline:
                doc.add_paragraph(f"Deadline: {deadline}")

        else:
            doc.add_paragraph(str(item), style="List Bullet")


def generate_docx(meeting):

    docx_path = OUTPUT_DIR / f"meeting_{meeting.id}.docx"

    doc = Document()

    title = doc.add_heading("MEETING MINUTES", level=1)
    title.runs[0].font.size = Pt(20)

    doc.add_paragraph(f"Meeting Title : {meeting.title}")
    doc.add_paragraph(f"Meeting Type : {meeting.meeting_type or 'N/A'}")
    doc.add_paragraph(f"Sentiment : {meeting.sentiment or 'N/A'}")

    summary = meeting.summary

    try:
        if isinstance(summary, str):
            summary = json.loads(summary)
    except:
        pass

    add_heading(doc, "Executive Summary")

    if isinstance(summary, dict):

        if summary.get("meeting_objective"):
            doc.add_paragraph(
                f"Meeting Objective:\n{summary['meeting_objective']}"
            )

        if summary.get("main_discussion"):
            doc.add_paragraph(
                f"Main Discussion:\n{summary['main_discussion']}"
            )

        conclusions = summary.get("important_conclusions", [])

        if conclusions:

            doc.add_paragraph("Important Conclusions:")

            if isinstance(conclusions, list):

                for item in conclusions:
                    doc.add_paragraph(item, style="List Bullet")

            else:
                doc.add_paragraph(conclusions)

        if summary.get("final_outcome"):
            doc.add_paragraph(
                f"Final Outcome:\n{summary['final_outcome']}"
            )

    else:
        doc.add_paragraph(str(summary))

    participants = load_json(meeting.participants)

    if participants:
        add_heading(doc, "Participants")
        add_list(doc, participants)

    key_points = load_json(meeting.key_points)

    if key_points:
        add_heading(doc, "Key Discussion Points")
        add_list(doc, key_points)

    decisions = load_json(meeting.decisions)

    if decisions:
        add_heading(doc, "Decisions Made")
        add_list(doc, decisions)

    action_items = load_json(meeting.action_items)

    if action_items:
        add_heading(doc, "Action Items")
        add_list(doc, action_items)

    risks = load_json(meeting.risks)

    if risks:
        add_heading(doc, "Risks & Concerns")
        add_list(doc, risks)

    questions = load_json(meeting.questions)

    if questions:
        add_heading(doc, "Open Questions")
        add_list(doc, questions)

    next_steps = load_json(meeting.next_steps)

    if next_steps:
        add_heading(doc, "Next Steps")
        add_list(doc, next_steps)

    if meeting.transcript:

        add_heading(doc, "Complete Meeting Transcript")
        doc.add_paragraph(meeting.transcript)

    doc.save(docx_path)

    return str(docx_path)